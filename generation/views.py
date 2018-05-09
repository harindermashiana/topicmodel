from django.views.generic.edit import FormView
from .form import FileFieldForm
from django.shortcuts import render,redirect

import docx
import PyPDF2
from gensim.parsing.preprocessing import STOPWORDS
documents = []
list1=[]
class FileFieldView(FormView):
    form_class = FileFieldForm
    template_name = 'base.html'  # Replace with your template.
   # success_url = 'moreorless' # Replace with your URL or reverse().
    files = []
    def post(self, request, *args, **kwargs):
        form_class = self.get_form_class()
        form = self.get_form(form_class)
        files=request.FILES.getlist('file_field')

        if form.is_valid():
            for f in files:

                writetotext(f)
            #return self.form_valid(form)
            return render(request,'moreorless.html',{
        })
        else:
            return self.form_invalid(form)
def writetotext(z):
  if (z not in list1):
    def fordocx(z):
        stringing = ""
        doc = docx.Document(z)
        for i in range(0, len(doc.paragraphs)):
            stringing = stringing + doc.paragraphs[i].text
        documents.append(stringing)

    def forpdf(z):
         pdfFileObj = z  # 'rb' for read binary mode
         pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
         x = pdfReader.numPages
         stringing = ""

         for t in range(0, x):
            texts = ""
            page = pdfReader.getPage(t)  # '9' is the page number
            stoplist = set('/n /u'.split())
            z = page.extractText()
            texts = [word for word in z.lower().split() if word not in stoplist]
            for xor in texts:
                stringing = stringing + " " + xor
        # documents.append(data)
         documents.append(stringing)
    filename=z.name
    if(filename.endswith('.pdf')):
        forpdf(z)
    if(filename.endswith('.docx')):
        fordocx(z)
    #fordocx(z)

    list1.append(z)

def modelx(request):
    return render(request, "moreorless.html", {})

def runmodel(request):
    from gensim import corpora, models, similarities
    import re
    import nltk

    from nltk.corpus import stopwords
    from nltk.stem.porter import PorterStemmer
    print(len(documents))
    stoplist = set(r'for a of the and to in \n'.split())
    texts = [[word for word in document.lower().split() if word not in STOPWORDS] for document in documents]
    from collections import defaultdict
    frequency = defaultdict(int)
    for text in texts:
        for token in text:
            frequency[token] += 1
    texts = [[token for token in text if frequency[token] > 1] for text in texts]
    dictionary = corpora.Dictionary(texts)
    dictionary.save('deerwester.dict')  # store the dictionary, for future reference
    corpus = [dictionary.doc2bow(text) for text in texts]
    corpora.MmCorpus.serialize('deerwester.mm', corpus)
    lsi = models.LsiModel(corpus, id2word=dictionary, num_topics=5)
    vb = lsi.show_topics(num_topics=-1, num_words=2, log=False, formatted=False)

    z=[]
    word=[]
    size=[]
    for t in vb:

        for a in t:
            if (type(a) == list):
                for h in a:
                    word.append(h[0])
                    size.append(h[1])
                word.append(' ')
                size.append(' ')

    x=2
    l=len(word)
    word.insert(0, 'Topic 1')
    for i in range(0,l-1):

        if(word[i]==' '):
            word[i]="Topic "+str(x)
            x=x+1
    print(word)
    size.insert(0, 'Topic 1')
    x=2
    for i in range(0, l - 1):

        if (size[i] == ' '):
            size[i] = "Topic " + str(x)
            x = x + 1
    print(size)
    return render(request, "table.html", {'word':word,
                                          'size':size})
